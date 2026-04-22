from __future__ import annotations

import hashlib
import re
import shutil
from datetime import datetime
from pathlib import Path
from typing import Any, List, Optional

from docx import Document
from reportlab.lib.pagesizes import A4
from reportlab.pdfgen import canvas

from .session_state import (
    ERROR_CATEGORY_FORMAT_UNSUPPORTED,
    ERROR_CATEGORY_PATH_VIOLATION,
    SCOPED_PATH_DENIED_TEXT,
    assert_access_granted_and_scoped,
    record_audit_event,
)


def sha256_text(text: str) -> str:
    return hashlib.sha256(text.encode("utf-8")).hexdigest()


def sha256_bytes(payload: bytes) -> str:
    return hashlib.sha256(payload).hexdigest()


def extract_numeric_values(payload: Any) -> List[float]:
    if isinstance(payload, list):
        if all(isinstance(item, (int, float)) for item in payload):
            return [float(item) for item in payload]

        values: List[float] = []
        for item in payload:
            if isinstance(item, dict):
                for value in item.values():
                    if isinstance(value, (int, float)):
                        values.append(float(value))
            elif isinstance(item, (int, float)):
                values.append(float(item))
        return values

    if isinstance(payload, dict):
        return [float(value) for value in payload.values() if isinstance(value, (int, float))]

    return []


def safe_report_path(title: str, report_dir: Path) -> Path:
    safe_title = re.sub(r"[^a-zA-Z0-9\u4e00-\u9fff_-]+", "-", title).strip("-_")
    safe_title = safe_title or "untitled-report"
    filename = f"{datetime.now().strftime('%Y-%m-%d')}-{safe_title}.md"

    target_path = (report_dir / filename).resolve()
    if report_dir not in target_path.parents and target_path != report_dir:
        raise ValueError("非法文件路径，已阻止目录穿越。")

    return target_path


def make_safe_stem(stem: str) -> str:
    safe_stem = re.sub(r"[^a-zA-Z0-9\u4e00-\u9fff_-]+", "-", stem).strip("-_")
    return safe_stem or "untitled-report"


def resolve_report_output_path(folder: str, title: str, format_name: str, filename: Optional[str]) -> Path:
    report_dir = assert_access_granted_and_scoped(folder)
    extension = f".{format_name}"

    if filename and filename.strip():
        raw_name = filename.strip()
        raw_path = Path(raw_name)
        if raw_path.name != raw_name or raw_path.parent != Path("."):
            raise PermissionError(SCOPED_PATH_DENIED_TEXT)
        stem = make_safe_stem(raw_path.stem)
        target_name = f"{stem}{extension}"
    else:
        date_prefix = datetime.now().strftime("%Y-%m-%d")
        target_name = f"{date_prefix}-{make_safe_stem(title)}{extension}"

    target_path = (report_dir / target_name).resolve()
    if report_dir not in target_path.parents:
        raise PermissionError(SCOPED_PATH_DENIED_TEXT)
    return target_path


def iter_content_lines(content: str) -> List[str]:
    cleaned = content.replace("\r\n", "\n").replace("\r", "\n").strip()
    return cleaned.split("\n") if cleaned else []


def write_docx_report(path: Path, title: str, content: str) -> None:
    doc = Document()
    doc.add_heading(title.strip(), level=1)
    for line in iter_content_lines(content):
        stripped = line.strip()
        if not stripped:
            doc.add_paragraph("")
            continue
        if stripped.startswith("#"):
            heading_text = stripped.lstrip("#").strip() or " "
            doc.add_heading(heading_text, level=2)
        else:
            doc.add_paragraph(stripped)
    doc.save(path)


def write_pdf_report(path: Path, title: str, content: str) -> None:
    pdf = canvas.Canvas(str(path), pagesize=A4)
    width, height = A4
    margin_x = 50
    y = height - 60

    pdf.setFont("Helvetica-Bold", 16)
    pdf.drawString(margin_x, y, title.strip() or "Untitled Report")
    y -= 30

    pdf.setFont("Helvetica", 11)
    for line in iter_content_lines(content):
        stripped = line.strip()
        if y < 50:
            pdf.showPage()
            pdf.setFont("Helvetica", 11)
            y = height - 60
        if stripped.startswith("#"):
            text = stripped.lstrip("#").strip() or " "
            pdf.setFont("Helvetica-Bold", 12)
            pdf.drawString(margin_x, y, text)
            pdf.setFont("Helvetica", 11)
        else:
            pdf.drawString(margin_x, y, stripped)
        y -= 18

    pdf.save()


def resolve_scoped_path(allowed_folder: str, filename: str) -> Path:
    report_dir = assert_access_granted_and_scoped(allowed_folder)
    filename_text = filename.strip()
    if not filename_text:
        raise ValueError("filename 不能为空。")

    candidate = Path(filename_text)
    if candidate.name != filename_text or candidate.parent != Path("."):
        record_audit_event(
            operation="path_resolution",
            target_file=filename_text,
            allowed_folder=allowed_folder.strip(),
            authorization_state="authorized",
            decision="deny",
            error_category=ERROR_CATEGORY_PATH_VIOLATION,
            details={"reason": "invalid_filename"},
        )
        raise PermissionError(SCOPED_PATH_DENIED_TEXT)

    target_path = (report_dir / candidate.name).resolve()
    if report_dir not in target_path.parents:
        record_audit_event(
            operation="path_resolution",
            target_file=filename_text,
            allowed_folder=str(report_dir),
            authorization_state="authorized",
            decision="deny",
            error_category=ERROR_CATEGORY_PATH_VIOLATION,
            details={"reason": "path_out_of_scope"},
        )
        raise PermissionError(SCOPED_PATH_DENIED_TEXT)
    record_audit_event(
        operation="path_resolution",
        target_file=target_path.name,
        allowed_folder=str(report_dir),
        authorization_state="authorized",
        decision="allow",
    )
    return target_path


def ensure_supported_read_extension(path: Path) -> str:
    ext = path.suffix.lower()
    if ext not in {".md", ".docx"}:
        if ext == ".pdf":
            record_audit_event(
                operation="read_report",
                target_file=path.name,
                allowed_folder=str(path.parent),
                authorization_state="authorized",
                decision="deny",
                error_category=ERROR_CATEGORY_FORMAT_UNSUPPORTED,
            )
            raise ValueError("MVP 暂不支持读取 PDF 正文；仅支持 .md/.docx。")
        record_audit_event(
            operation="read_report",
            target_file=path.name,
            allowed_folder=str(path.parent),
            authorization_state="authorized",
            decision="deny",
            error_category=ERROR_CATEGORY_FORMAT_UNSUPPORTED,
        )
        raise ValueError("仅支持 .md/.docx 文件。")
    return ext


def ensure_supported_edit_extension(path: Path) -> str:
    ext = path.suffix.lower()
    if ext == ".pdf":
        record_audit_event(
            operation="edit_report",
            target_file=path.name,
            allowed_folder=str(path.parent),
            authorization_state="authorized",
            decision="deny",
            error_category=ERROR_CATEGORY_FORMAT_UNSUPPORTED,
        )
        raise ValueError("MVP 暂不支持编辑 PDF；请改用 .md/.docx。")
    if ext not in {".md", ".docx"}:
        record_audit_event(
            operation="edit_report",
            target_file=path.name,
            allowed_folder=str(path.parent),
            authorization_state="authorized",
            decision="deny",
            error_category=ERROR_CATEGORY_FORMAT_UNSUPPORTED,
        )
        raise ValueError("仅支持编辑 .md/.docx 文件。")
    return ext


def make_backup(path: Path) -> Path:
    timestamp = datetime.now().strftime("%Y%m%d%H%M%S")
    backup_path = path.with_name(f"{path.name}.bak.{timestamp}")
    shutil.copy2(path, backup_path)
    return backup_path


def parse_replace_instruction(instruction: str) -> tuple[str, str]:
    text = instruction.strip()
    if "\n---\n" in text:
        head, body = text.split("\n---\n", 1)
    else:
        lines = text.splitlines()
        if len(lines) < 2:
            raise ValueError("replace_section 模式需要 section 标题和新内容。")
        head, body = lines[0], "\n".join(lines[1:])

    section = re.sub(r"^section\s*:\s*", "", head.strip(), flags=re.IGNORECASE).strip().lstrip("#").strip()
    if not section:
        raise ValueError("replace_section 模式缺少 section 标题。")
    return section, body.strip()


def extract_md_headings(lines: List[str]) -> List[tuple[int, int, str]]:
    headings: List[tuple[int, int, str]] = []
    for idx, line in enumerate(lines):
        match = re.match(r"^(#{1,6})\s+(.*)$", line.strip())
        if not match:
            continue
        level = len(match.group(1))
        title = match.group(2).strip()
        headings.append((idx, level, title))
    return headings


def edit_markdown_content(original: str, instruction: str, mode: str) -> tuple[str, List[str], int]:
    normalized = original.replace("\r\n", "\n").replace("\r", "\n")
    lines = normalized.split("\n")
    changed_sections: List[str] = []

    if mode == "append":
        append_text = instruction.strip()
        if not append_text:
            raise ValueError("append 模式下 instruction 不能为空。")
        merged = normalized.rstrip("\n")
        if merged:
            merged = f"{merged}\n\n{append_text}\n"
        else:
            merged = f"{append_text}\n"
        return merged, ["__appended__"], len(append_text.splitlines())

    if mode == "rewrite":
        rewritten = instruction.strip()
        if not rewritten:
            raise ValueError("rewrite 模式下 instruction 不能为空。")
        return f"{rewritten}\n", ["__all__"], len(rewritten.splitlines())

    section_name, replacement = parse_replace_instruction(instruction)
    headings = extract_md_headings(lines)
    target_index = None
    target_level = None
    for idx, level, title in headings:
        if title.lower() == section_name.lower():
            target_index = idx
            target_level = level
            break
    if target_index is None or target_level is None:
        raise ValueError(f"未找到 Markdown 节：{section_name}")

    end_index = len(lines)
    for idx, level, _ in headings:
        if idx > target_index and level <= target_level:
            end_index = idx
            break

    heading_line = lines[target_index]
    replacement_block = [heading_line]
    replacement_lines = replacement.splitlines()
    if replacement_lines:
        replacement_block.extend(replacement_lines)
    new_lines = lines[:target_index] + replacement_block + lines[end_index:]
    changed_sections.append(section_name)
    return "\n".join(new_lines).rstrip("\n") + "\n", changed_sections, len(replacement_lines)
