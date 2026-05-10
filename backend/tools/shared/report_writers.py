from pathlib import Path
from typing import List

from docx import Document
from reportlab.lib.pagesizes import A4
from reportlab.pdfgen import canvas


def iter_content_lines(content: str) -> List[str]:
    cleaned = content.replace("\r\n", "\n").replace("\r", "\n").strip()
    return cleaned.split("\n") if cleaned else []


def write_docx_report(path: Path, title: str, content: str) -> None:
    doc = Document()
    doc.add_heading(title.strip(), level=1)
    for line in iter_content_lines(content):
        doc.add_paragraph(line.strip())
    doc.save(path)


def write_pdf_report(path: Path, title: str, content: str) -> None:
    pdf = canvas.Canvas(str(path), pagesize=A4)
    _, height = A4
    y = height - 60
    pdf.setFont("Helvetica-Bold", 16)
    pdf.drawString(50, y, title.strip() or "Untitled Report")
    y -= 30
    pdf.setFont("Helvetica", 11)
    for line in iter_content_lines(content):
        if y < 50:
            pdf.showPage(); pdf.setFont("Helvetica", 11); y = height - 60
        pdf.drawString(50, y, line.strip())
        y -= 18
    pdf.save()
