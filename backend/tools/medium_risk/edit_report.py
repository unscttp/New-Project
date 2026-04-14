import json
from typing import Callable, Literal

from docx import Document


def edit_report(
    file_name: str,
    folder: str,
    instruction: str,
    mode: Literal["append", "replace_section", "rewrite"],
    *,
    resolve_scoped_path: Callable[[str, str], object],
    scoped_path_denied_text: str,
    ensure_supported_edit_extension: Callable[[object], str],
    make_backup: Callable[[object], object],
    sha256_text: Callable[[str], str],
    sha256_bytes: Callable[[bytes], str],
    edit_markdown_content: Callable[[str, str, str], tuple[str, list[str], int]],
    parse_replace_instruction: Callable[[str], tuple[str, str]],
    record_audit_event: Callable[..., None],
) -> str:
    target_path = resolve_scoped_path(folder, file_name)
    if not target_path.exists():
        raise FileNotFoundError(f"文件不存在：{target_path.name}")
    if not target_path.is_file():
        raise PermissionError(scoped_path_denied_text)

    ext = ensure_supported_edit_extension(target_path)
    backup_path = make_backup(target_path)

    if ext == ".md":
        original = target_path.read_text(encoding="utf-8")
        before_hash = sha256_text(original)
        new_content, sections, touched_lines = edit_markdown_content(original, instruction, mode)
        target_path.write_text(new_content, encoding="utf-8")
        after_hash = sha256_text(new_content)
        summary = {
            "file_name": target_path.name,
            "format": "md",
            "mode": mode,
            "changed_sections": sections,
            "line_count_before": len(original.splitlines()),
            "line_count_after": len(new_content.splitlines()),
            "touched_line_count": touched_lines,
            "backup_path": str(backup_path),
            "checksum_before": before_hash,
            "checksum_after": after_hash,
        }
        record_audit_event(
            operation="edit_report",
            target_file=target_path.name,
            allowed_folder=str(target_path.parent),
            authorization_state="authorized",
            decision="allow",
            details={
                "mode": mode,
                "changed_sections": sections,
                "checksum_before": before_hash,
                "checksum_after": after_hash,
            },
        )
        return json.dumps(summary, ensure_ascii=False, indent=2)

    doc = Document(target_path)
    before_hash = sha256_bytes(target_path.read_bytes())
    changed_sections: list[str] = []
    paragraph_count_before = len(doc.paragraphs)
    touched_paragraphs = 0

    if mode == "append":
        append_text = instruction.strip()
        if not append_text:
            raise ValueError("append 模式下 instruction 不能为空。")
        for line in append_text.splitlines():
            doc.add_paragraph(line)
            touched_paragraphs += 1
        changed_sections = ["__appended__"]
    elif mode == "rewrite":
        for _ in range(len(doc.paragraphs)):
            p = doc.paragraphs[0]._element
            p.getparent().remove(p)
        for line in instruction.strip().splitlines():
            doc.add_paragraph(line)
            touched_paragraphs += 1
        changed_sections = ["__all__"]
    else:
        section_name, replacement = parse_replace_instruction(instruction)
        replacement_lines = replacement.splitlines()
        start_idx = None
        for idx, paragraph in enumerate(doc.paragraphs):
            if paragraph.text.strip().lower() == section_name.lower():
                start_idx = idx
                break
        if start_idx is None:
            raise ValueError(f"未找到 DOCX 段落标题：{section_name}")
        end_idx = len(doc.paragraphs)
        for idx in range(start_idx + 1, len(doc.paragraphs)):
            if doc.paragraphs[idx].style and str(doc.paragraphs[idx].style.name).lower().startswith("heading"):
                end_idx = idx
                break
        anchor = doc.paragraphs[start_idx]._element
        for _ in range(end_idx - start_idx - 1):
            nxt = anchor.getnext()
            if nxt is not None:
                nxt.getparent().remove(nxt)
                touched_paragraphs += 1
        for line in replacement_lines:
            new_para = doc.add_paragraph(line)
            anchor.addnext(new_para._element)
            anchor = new_para._element
            touched_paragraphs += 1
        changed_sections = [section_name]

    doc.save(target_path)
    after_hash = sha256_bytes(target_path.read_bytes())
    summary = {
        "file_name": target_path.name,
        "format": "docx",
        "mode": mode,
        "changed_sections": changed_sections,
        "paragraph_count_before": paragraph_count_before,
        "paragraph_count_after": len(Document(target_path).paragraphs),
        "touched_paragraph_count": touched_paragraphs,
        "backup_path": str(backup_path),
        "checksum_before": before_hash,
        "checksum_after": after_hash,
    }
    record_audit_event(
        operation="edit_report",
        target_file=target_path.name,
        allowed_folder=str(target_path.parent),
        authorization_state="authorized",
        decision="allow",
        details={
            "mode": mode,
            "changed_sections": changed_sections,
            "checksum_before": before_hash,
            "checksum_after": after_hash,
        },
    )
    return json.dumps(summary, ensure_ascii=False, indent=2)
