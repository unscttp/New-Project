import json
import re
import shutil
import hashlib
import importlib
from datetime import datetime
from pathlib import Path
from typing import Any, Dict, List, Literal, Optional

from docx import Document
from reportlab.lib.pagesizes import A4
from reportlab.pdfgen import canvas

from .high_risk.delete_report_file import delete_report_file as delete_report_file_high_risk
from .low_risk.analyze_trend_data import analyze_trend_data as analyze_trend_data_low_risk
from .low_risk.confirm_report_folder_access import (
    confirm_report_folder_access as confirm_report_folder_access_low_risk,
)
from .low_risk.creation import save_report_file as save_report_file_low_risk
from .low_risk.generate_markdown_report import (
    generate_markdown_report as generate_markdown_report_low_risk,
)
from .low_risk.list_report_files import list_report_files as list_report_files_low_risk
from .low_risk.read_report_file import read_report_file as read_report_file_low_risk
from .low_risk.request_report_folder_access import (
    request_report_folder_access as request_report_folder_access_low_risk,
)
from .low_risk.save_report import save_report as save_report_low_risk
from .low_risk.search_internet import search_internet as search_internet_low_risk
from .low_risk.select_tool_risk_level import (
    select_tool_risk_level as select_tool_risk_level_low_risk,
)
from .low_risk.extract_keywords import extract_keywords as extract_keywords_low_risk
from .medium_risk.editing import edit_report_file as edit_report_file_medium_risk
from .medium_risk.edit_report import edit_report as edit_report_medium_risk
from .medium_risk.read_report import read_report as read_report_medium_risk
from .medium_risk.redact_sensitive_text import redact_sensitive_text as redact_sensitive_text_medium_risk
from .risk_control import (
    assert_tool_access,
    set_active_risk_level,
    set_active_session as set_active_risk_session,
)


BASE_DIR = Path(__file__).resolve().parent
UNAUTHORIZED_FILE_ACCESS_TEXT = "未授权文件访问，请先确认目录和权限。"
SCOPED_PATH_DENIED_TEXT = "目标文件不在授权目录内，操作已拒绝。"
ALLOWED_REPORT_EXTENSIONS = {".md", ".docx", ".pdf"}
ACTIVE_SESSION_ID: str = "default"
SESSION_PERMISSION_STATE: Dict[str, Dict[str, Optional[str] | bool]] = {}
SESSION_AUDIT_LOGS: Dict[str, List[Dict[str, Any]]] = {}
ERROR_CATEGORY_PERMISSION_DENIED = "permission_denied"
ERROR_CATEGORY_PATH_VIOLATION = "path_violation"
ERROR_CATEGORY_FORMAT_UNSUPPORTED = "format_unsupported"
ERROR_CATEGORY_IO_FAILURE = "io_failure"


def set_active_session(session_id: str, initial_state: Optional[Dict[str, Any]] = None) -> None:
    global ACTIVE_SESSION_ID
    ACTIVE_SESSION_ID = (session_id or "default").strip() or "default"
    if initial_state is not None:
        SESSION_PERMISSION_STATE[ACTIVE_SESSION_ID] = {
            "file_access_granted": bool(initial_state.get("file_access_granted", False)),
            "allowed_report_folder": initial_state.get("allowed_report_folder"),
        }
    elif ACTIVE_SESSION_ID not in SESSION_PERMISSION_STATE:
        SESSION_PERMISSION_STATE[ACTIVE_SESSION_ID] = {
            "file_access_granted": False,
            "allowed_report_folder": None,
        }
    SESSION_AUDIT_LOGS.setdefault(ACTIVE_SESSION_ID, [])
    set_active_risk_session(ACTIVE_SESSION_ID)


def get_active_permission_state() -> Dict[str, Optional[str] | bool]:
    return SESSION_PERMISSION_STATE.get(
        ACTIVE_SESSION_ID,
        {"file_access_granted": False, "allowed_report_folder": None},
    )


def get_active_audit_entries() -> List[Dict[str, Any]]:
    return list(SESSION_AUDIT_LOGS.get(ACTIVE_SESSION_ID, []))


def _sha256_text(text: str) -> str:
    return hashlib.sha256(text.encode("utf-8")).hexdigest()


def _sha256_bytes(payload: bytes) -> str:
    return hashlib.sha256(payload).hexdigest()


def record_audit_event(
    operation: str,
    allowed_folder: Optional[str],
    authorization_state: str,
    decision: str,
    target_file: Optional[str] = None,
    error_category: Optional[str] = None,
    details: Optional[Dict[str, Any]] = None,
) -> None:
    event: Dict[str, Any] = {
        "timestamp": datetime.utcnow().isoformat(timespec="seconds") + "Z",
        "operation": operation,
        "target_file": target_file,
        "allowed_folder": allowed_folder,
        "authorization_state": authorization_state,
        "decision": decision,
    }
    if error_category:
        event["error_category"] = error_category
    if details:
        event["details"] = details
    SESSION_AUDIT_LOGS.setdefault(ACTIVE_SESSION_ID, []).append(event)


def _extract_numeric_values(payload: Any) -> List[float]:
    if isinstance(payload, list):
        if all(isinstance(item, (int, float)) for item in payload):
            return [float(item) for item in payload]

        values: List[float] = []
        for item in payload:
            if isinstance(item, dict):
                for value in item.values():
                    if isinstance(value, (int, float)):
                        values.append(float(value))
            elif isinstance(item, (int, float)):
                values.append(float(item))
        return values

    if isinstance(payload, dict):
        return [float(value) for value in payload.values() if isinstance(value, (int, float))]

    return []


def search_internet(query: str) -> str:
    return search_internet_low_risk(query)


def analyze_trend_data(data_json: str) -> str:
    return analyze_trend_data_low_risk(data_json, extract_numeric_values=_extract_numeric_values)


def _safe_report_path(title: str, report_dir: Path) -> Path:
    safe_title = re.sub(r"[^a-zA-Z0-9\u4e00-\u9fff_-]+", "-", title).strip("-_")
    safe_title = safe_title or "untitled-report"
    filename = f"{datetime.now().strftime('%Y-%m-%d')}-{safe_title}.md"

    target_path = (report_dir / filename).resolve()
    if report_dir not in target_path.parents and target_path != report_dir:
        raise ValueError("非法文件路径，已阻止目录穿越。")

    return target_path


def _make_safe_stem(stem: str) -> str:
    safe_stem = re.sub(r"[^a-zA-Z0-9\u4e00-\u9fff_-]+", "-", stem).strip("-_")
    return safe_stem or "untitled-report"


def _resolve_report_output_path(folder: str, title: str, format_name: str, filename: Optional[str]) -> Path:
    report_dir = assert_access_granted_and_scoped(folder)
    extension = f".{format_name}"

    if filename and filename.strip():
        raw_name = filename.strip()
        raw_path = Path(raw_name)
        if raw_path.name != raw_name or raw_path.parent != Path("."):
            raise PermissionError(SCOPED_PATH_DENIED_TEXT)
        stem = _make_safe_stem(raw_path.stem)
        target_name = f"{stem}{extension}"
    else:
        date_prefix = datetime.now().strftime("%Y-%m-%d")
        target_name = f"{date_prefix}-{_make_safe_stem(title)}{extension}"

    target_path = (report_dir / target_name).resolve()
    if report_dir not in target_path.parents:
        raise PermissionError(SCOPED_PATH_DENIED_TEXT)
    return target_path


def _iter_content_lines(content: str) -> List[str]:
    cleaned = content.replace("\r\n", "\n").replace("\r", "\n").strip()
    return cleaned.split("\n") if cleaned else []


def _write_docx_report(path: Path, title: str, content: str) -> None:
    doc = Document()
    doc.add_heading(title.strip(), level=1)
    for line in _iter_content_lines(content):
        stripped = line.strip()
        if not stripped:
            doc.add_paragraph("")
            continue
        if stripped.startswith("#"):
            heading_text = stripped.lstrip("#").strip() or " "
            doc.add_heading(heading_text, level=2)
        else:
            doc.add_paragraph(stripped)
    doc.save(path)


def _write_pdf_report(path: Path, title: str, content: str) -> None:
    pdf = canvas.Canvas(str(path), pagesize=A4)
    width, height = A4
    margin_x = 50
    y = height - 60

    pdf.setFont("Helvetica-Bold", 16)
    pdf.drawString(margin_x, y, title.strip() or "Untitled Report")
    y -= 30

    pdf.setFont("Helvetica", 11)
    for line in _iter_content_lines(content):
        stripped = line.strip()
        if y < 50:
            pdf.showPage()
            pdf.setFont("Helvetica", 11)
            y = height - 60
        if stripped.startswith("#"):
            text = stripped.lstrip("#").strip() or " "
            pdf.setFont("Helvetica-Bold", 12)
            pdf.drawString(margin_x, y, text)
            pdf.setFont("Helvetica", 11)
        else:
            pdf.drawString(margin_x, y, stripped)
        y -= 18

    pdf.save()


def assert_access_granted_and_scoped(allowed_folder: str) -> Path:
    state = get_active_permission_state()
    granted = bool(state.get("file_access_granted"))
    scoped_folder = state.get("allowed_report_folder")
    allowed_folder_text = allowed_folder.strip()
    if not granted or not scoped_folder or not str(scoped_folder).strip():
        record_audit_event(
            operation="permission_check",
            allowed_folder=allowed_folder_text or str(scoped_folder or ""),
            authorization_state="unauthorized",
            decision="deny",
            error_category=ERROR_CATEGORY_PERMISSION_DENIED,
        )
        raise PermissionError(UNAUTHORIZED_FILE_ACCESS_TEXT)
    if not allowed_folder_text:
        raise ValueError("allowed_folder 不能为空。")

    state_dir = Path(str(scoped_folder)).expanduser().resolve()
    request_dir = Path(allowed_folder_text).expanduser().resolve()
    if request_dir != state_dir:
        record_audit_event(
            operation="permission_check",
            allowed_folder=allowed_folder_text,
            authorization_state="authorized",
            decision="deny",
            error_category=ERROR_CATEGORY_PATH_VIOLATION,
            details={"reason": "folder_not_in_scope"},
        )
        raise PermissionError(UNAUTHORIZED_FILE_ACCESS_TEXT)

    report_dir = state_dir
    report_dir.mkdir(parents=True, exist_ok=True)
    record_audit_event(
        operation="permission_check",
        allowed_folder=str(report_dir),
        authorization_state="authorized",
        decision="allow",
    )
    return report_dir


def resolve_scoped_path(allowed_folder: str, filename: str) -> Path:
    report_dir = assert_access_granted_and_scoped(allowed_folder)
    filename_text = filename.strip()
    if not filename_text:
        raise ValueError("filename 不能为空。")

    candidate = Path(filename_text)
    if candidate.name != filename_text or candidate.parent != Path("."):
        record_audit_event(
            operation="path_resolution",
            target_file=filename_text,
            allowed_folder=allowed_folder.strip(),
            authorization_state="authorized",
            decision="deny",
            error_category=ERROR_CATEGORY_PATH_VIOLATION,
            details={"reason": "invalid_filename"},
        )
        raise PermissionError(SCOPED_PATH_DENIED_TEXT)

    target_path = (report_dir / candidate.name).resolve()
    if report_dir not in target_path.parents:
        record_audit_event(
            operation="path_resolution",
            target_file=filename_text,
            allowed_folder=str(report_dir),
            authorization_state="authorized",
            decision="deny",
            error_category=ERROR_CATEGORY_PATH_VIOLATION,
            details={"reason": "path_out_of_scope"},
        )
        raise PermissionError(SCOPED_PATH_DENIED_TEXT)
    record_audit_event(
        operation="path_resolution",
        target_file=target_path.name,
        allowed_folder=str(report_dir),
        authorization_state="authorized",
        decision="allow",
    )
    return target_path


def _ensure_supported_read_extension(path: Path) -> str:
    ext = path.suffix.lower()
    if ext not in {".md", ".docx"}:
        if ext == ".pdf":
            record_audit_event(
                operation="read_report",
                target_file=path.name,
                allowed_folder=str(path.parent),
                authorization_state="authorized",
                decision="deny",
                error_category=ERROR_CATEGORY_FORMAT_UNSUPPORTED,
            )
            raise ValueError("MVP 暂不支持读取 PDF 正文；仅支持 .md/.docx。")
        record_audit_event(
            operation="read_report",
            target_file=path.name,
            allowed_folder=str(path.parent),
            authorization_state="authorized",
            decision="deny",
            error_category=ERROR_CATEGORY_FORMAT_UNSUPPORTED,
        )
        raise ValueError("仅支持 .md/.docx 文件。")
    return ext


def _ensure_supported_edit_extension(path: Path) -> str:
    ext = path.suffix.lower()
    if ext == ".pdf":
        record_audit_event(
            operation="edit_report",
            target_file=path.name,
            allowed_folder=str(path.parent),
            authorization_state="authorized",
            decision="deny",
            error_category=ERROR_CATEGORY_FORMAT_UNSUPPORTED,
        )
        raise ValueError("MVP 暂不支持编辑 PDF；请改用 .md/.docx。")
    if ext not in {".md", ".docx"}:
        record_audit_event(
            operation="edit_report",
            target_file=path.name,
            allowed_folder=str(path.parent),
            authorization_state="authorized",
            decision="deny",
            error_category=ERROR_CATEGORY_FORMAT_UNSUPPORTED,
        )
        raise ValueError("仅支持编辑 .md/.docx 文件。")
    return ext


def _make_backup(path: Path) -> Path:
    timestamp = datetime.now().strftime("%Y%m%d%H%M%S")
    backup_path = path.with_name(f"{path.name}.bak.{timestamp}")
    shutil.copy2(path, backup_path)
    return backup_path


def _parse_replace_instruction(instruction: str) -> tuple[str, str]:
    text = instruction.strip()
    if "\n---\n" in text:
        head, body = text.split("\n---\n", 1)
    else:
        lines = text.splitlines()
        if len(lines) < 2:
            raise ValueError("replace_section 模式需要 section 标题和新内容。")
        head, body = lines[0], "\n".join(lines[1:])

    section = re.sub(r"^section\s*:\s*", "", head.strip(), flags=re.IGNORECASE).strip().lstrip("#").strip()
    if not section:
        raise ValueError("replace_section 模式缺少 section 标题。")
    return section, body.strip()


def _extract_md_headings(lines: List[str]) -> List[tuple[int, int, str]]:
    headings: List[tuple[int, int, str]] = []
    for idx, line in enumerate(lines):
        match = re.match(r"^(#{1,6})\s+(.*)$", line.strip())
        if not match:
            continue
        level = len(match.group(1))
        title = match.group(2).strip()
        headings.append((idx, level, title))
    return headings


def _edit_markdown_content(original: str, instruction: str, mode: str) -> tuple[str, List[str], int]:
    normalized = original.replace("\r\n", "\n").replace("\r", "\n")
    lines = normalized.split("\n")
    changed_sections: List[str] = []

    if mode == "append":
        append_text = instruction.strip()
        if not append_text:
            raise ValueError("append 模式下 instruction 不能为空。")
        merged = normalized.rstrip("\n")
        if merged:
            merged = f"{merged}\n\n{append_text}\n"
        else:
            merged = f"{append_text}\n"
        return merged, ["__appended__"], len(append_text.splitlines())

    if mode == "rewrite":
        rewritten = instruction.strip()
        if not rewritten:
            raise ValueError("rewrite 模式下 instruction 不能为空。")
        return f"{rewritten}\n", ["__all__"], len(rewritten.splitlines())

    section_name, replacement = _parse_replace_instruction(instruction)
    headings = _extract_md_headings(lines)
    target_index = None
    target_level = None
    for idx, level, title in headings:
        if title.lower() == section_name.lower():
            target_index = idx
            target_level = level
            break
    if target_index is None or target_level is None:
        raise ValueError(f"未找到 Markdown 节：{section_name}")

    end_index = len(lines)
    for idx, level, _ in headings:
        if idx > target_index and level <= target_level:
            end_index = idx
            break

    heading_line = lines[target_index]
    replacement_block = [heading_line]
    replacement_lines = replacement.splitlines()
    if replacement_lines:
        replacement_block.extend(replacement_lines)
    new_lines = lines[:target_index] + replacement_block + lines[end_index:]
    changed_sections.append(section_name)
    return "\n".join(new_lines).rstrip("\n") + "\n", changed_sections, len(replacement_lines)


def request_report_folder_access(purpose: str, folder: str) -> str:
    return request_report_folder_access_low_risk(purpose, folder)


def confirm_report_folder_access(granted: bool, folder: str) -> str:
    return confirm_report_folder_access_low_risk(
        granted,
        folder,
        session_permission_state=SESSION_PERMISSION_STATE,
        active_session_id=ACTIVE_SESSION_ID,
        record_audit_event=record_audit_event,
        unauthorized_file_access_text=UNAUTHORIZED_FILE_ACCESS_TEXT,
        error_category_permission_denied=ERROR_CATEGORY_PERMISSION_DENIED,
    )


def generate_markdown_report(title: str, content: str) -> str:
    return generate_markdown_report_low_risk(
        title,
        content=content,
        get_active_permission_state=get_active_permission_state,
        assert_access_granted_and_scoped=assert_access_granted_and_scoped,
        safe_report_path=_safe_report_path,
    )







def _build_tool_risk_levels(tool_specs: List[Dict[str, Any]]) -> Dict[str, str]:
    levels: Dict[str, str] = {}
    for spec in tool_specs:
        name = str(spec.get("name", "")).strip()
        risk_level = str(spec.get("risk_level", "")).strip().lower()
        if not name:
            raise ValueError("tool_registry.json 中存在缺少 name 的工具配置")
        if risk_level not in {"low", "medium", "high"}:
            raise ValueError(f"工具 {name} 的 risk_level 非法: {risk_level}")
        levels[name] = risk_level
    return levels


def _enforce_tool_risk(tool_name: str) -> None:
    required = TOOL_RISK_LEVELS.get(tool_name, "high")
    assert_tool_access(required, tool_name)


def select_tool_risk_level(risk_level: Literal["low", "medium", "high"]) -> str:
    return select_tool_risk_level_low_risk(risk_level, set_active_risk_level=set_active_risk_level)

def extract_keywords(text: str, top_k: int = 8) -> str:
    _enforce_tool_risk("extract_keywords")
    return extract_keywords_low_risk(text=text, top_k=top_k)


def redact_sensitive_text(content: str) -> str:
    _enforce_tool_risk("redact_sensitive_text")
    return redact_sensitive_text_medium_risk(content=content)


def delete_report_file(allowed_folder: str, filename: str) -> str:
    _enforce_tool_risk("delete_report_file")
    return delete_report_file_high_risk(
        allowed_folder=allowed_folder,
        filename=filename,
        resolve_scoped_path=resolve_scoped_path,
        record_audit_event=record_audit_event,
    )

def save_report(
    title: str,
    content: str,
    format: Literal["md", "docx", "pdf"],
    folder: str,
    filename: Optional[str] = None,
) -> str:
    _enforce_tool_risk("save_report")
    return save_report_low_risk(
        title=title,
        content=content,
        format=format,
        folder=folder,
        filename=filename,
        resolve_report_output_path=_resolve_report_output_path,
        write_docx_report=_write_docx_report,
        write_pdf_report=_write_pdf_report,
    )


def save_report_file(allowed_folder: str, filename: str, content: str) -> str:
    _enforce_tool_risk("save_report_file")
    return save_report_file_low_risk(
        allowed_folder,
        filename,
        content,
        resolve_scoped_path=resolve_scoped_path,
        sha256_text=_sha256_text,
        record_audit_event=record_audit_event,
    )


def read_report_file(allowed_folder: str, filename: str) -> str:
    _enforce_tool_risk("read_report_file")
    return read_report_file_low_risk(
        allowed_folder=allowed_folder,
        filename=filename,
        resolve_scoped_path=resolve_scoped_path,
        scoped_path_denied_text=SCOPED_PATH_DENIED_TEXT,
    )


def edit_report_file(allowed_folder: str, filename: str, content: str) -> str:
    _enforce_tool_risk("edit_report_file")
    return edit_report_file_medium_risk(
        allowed_folder,
        filename,
        content,
        resolve_scoped_path=resolve_scoped_path,
        scoped_path_denied_text=SCOPED_PATH_DENIED_TEXT,
        sha256_text=_sha256_text,
        record_audit_event=record_audit_event,
    )


def read_report(file_name: str, folder: str) -> str:
    _enforce_tool_risk("read_report")
    return read_report_medium_risk(
        file_name=file_name,
        folder=folder,
        resolve_scoped_path=resolve_scoped_path,
        scoped_path_denied_text=SCOPED_PATH_DENIED_TEXT,
        ensure_supported_read_extension=_ensure_supported_read_extension,
    )


def edit_report(file_name: str, folder: str, instruction: str, mode: Literal["append", "replace_section", "rewrite"]) -> str:
    _enforce_tool_risk("edit_report")
    return edit_report_medium_risk(
        file_name=file_name,
        folder=folder,
        instruction=instruction,
        mode=mode,
        resolve_scoped_path=resolve_scoped_path,
        scoped_path_denied_text=SCOPED_PATH_DENIED_TEXT,
        ensure_supported_edit_extension=_ensure_supported_edit_extension,
        make_backup=_make_backup,
        sha256_text=_sha256_text,
        sha256_bytes=_sha256_bytes,
        edit_markdown_content=_edit_markdown_content,
        parse_replace_instruction=_parse_replace_instruction,
        record_audit_event=record_audit_event,
    )


def list_report_files(allowed_folder: str) -> str:
    _enforce_tool_risk("list_report_files")
    return list_report_files_low_risk(
        allowed_folder=allowed_folder,
        assert_access_granted_and_scoped=assert_access_granted_and_scoped,
        allowed_report_extensions=ALLOWED_REPORT_EXTENSIONS,
    )


TOOL_REGISTRY_CONFIG_PATH = BASE_DIR / "tool_registry.json"


def _load_tool_registry_config() -> List[Dict[str, Any]]:
    raw = json.loads(TOOL_REGISTRY_CONFIG_PATH.read_text(encoding="utf-8"))
    tools = raw.get("tools", [])
    if not isinstance(tools, list):
        raise ValueError("tool_registry.json 的 tools 字段必须为列表")
    return tools


def _resolve_object(import_path: str) -> Any:
    if ":" not in import_path:
        raise ValueError(f"导入路径格式错误: {import_path}，应为 module.path:object_name")
    module_path, object_name = import_path.split(":", 1)
    module = importlib.import_module(module_path)
    return getattr(module, object_name)


def _build_tool_registry(tool_specs: List[Dict[str, Any]]) -> Dict[str, Any]:
    registry: Dict[str, Any] = {}
    for spec in tool_specs:
        name = spec["name"]
        registry[name] = _resolve_object(spec["callable_path"])
    return registry


def _build_openai_tools(tool_specs: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
    payload: List[Dict[str, Any]] = []
    for spec in tool_specs:
        args_model = _resolve_object(spec["args_model_path"])
        payload.append(
            {
                "type": "function",
                "function": {
                    "name": spec["name"],
                    "description": spec["description"],
                    "parameters": args_model.model_json_schema(),
                },
            }
        )
    return payload


_TOOL_SPECS = _load_tool_registry_config()
TOOL_RISK_LEVELS = _build_tool_risk_levels(_TOOL_SPECS)
TOOL_REGISTRY = _build_tool_registry(_TOOL_SPECS)
OPENAI_TOOLS: List[Dict[str, Any]] = _build_openai_tools(_TOOL_SPECS)

